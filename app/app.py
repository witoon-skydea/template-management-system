import os
import sys
from flask import Flask, render_template, redirect, url_for, flash, request, jsonify, session
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from werkzeug.security import check_password_hash
from passlib.hash import pbkdf2_sha256
import markdown
import json
import re
from docx import Document as DocxDocument
import bleach
from functools import wraps
import copy
import datetime
from sqlalchemy import and_, or_, desc

# Add parent directory to path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from sqlalchemy.orm import joinedload
from database.models import init_db, get_session, User, Template, InputBox, Document, DocumentInputValue, TemplateAssignment, Notification, Station, StationUser, StationUserRole, ChatChannel, ChatMessage, ChatChannelType

# Configure application
app = Flask(__name__, 
            static_folder=os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "static"),
            template_folder=os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "templates"))
app.config['SECRET_KEY'] = 'template-management-secret-key'

# Configure database
db_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "database", "template_management.db")
engine = init_db(db_path)

# Configure login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

@login_manager.user_loader
def load_user(user_id):
    db_session = get_session(engine)
    user = db_session.query(User).get(int(user_id))
    db_session.close()
    return user

# Admin required decorator
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_admin:
            flash('You do not have permission to access this page', 'danger')
            return redirect(url_for('dashboard'))
        return f(*args, **kwargs)
    return decorated_function

# Station master required decorator
def station_master_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated:
            flash('You must be logged in to access this page', 'danger')
            return redirect(url_for('login'))
        
        # If user is admin, allow access
        if current_user.is_admin:
            return f(*args, **kwargs)
            
        # Check if station_id is in the route parameters
        station_id = kwargs.get('station_id')
        if not station_id:
            flash('Station ID is required', 'danger')
            return redirect(url_for('dashboard'))
            
        # Check if user is a station master for this station
        db_session = get_session(engine)
        station_user = db_session.query(StationUser).filter(
            StationUser.station_id == station_id,
            StationUser.user_id == current_user.id,
            StationUser.role == StationUserRole.STATION_MASTER.value
        ).first()
        db_session.close()
        
        if not station_user:
            flash('You do not have permission to access this page', 'danger')
            return redirect(url_for('dashboard'))
            
        return f(*args, **kwargs)
    return decorated_function

# Helper functions
def extract_input_boxes(content):
    """
    Extract input boxes from markdown content using the *****input-box***** syntax
    Returns a list of box IDs
    """
    pattern = r'\*\*\*\*\*(.*?)\*\*\*\*\*'
    input_boxes = re.findall(pattern, content)
    return input_boxes

def markdown_to_html(md_content, input_boxes=None, input_values=None, document_id=None):
    """
    Convert markdown to HTML with enhanced extensions for better rendering
    
    If input_boxes and input_values are provided, the function will add
    special markup for input boxes to make them interactive in the UI
    """
    try:
        # Process input boxes if provided
        processed_content = md_content
        box_mappings = {}
        
        if input_boxes and document_id:
            # Create a map of input box placeholders to their IDs for later replacement
            for box in input_boxes:
                placeholder = f'*****{box.box_id}*****'
                value = ""
                
                # Find the value for this box if input_values provided
                if input_values:
                    for val in input_values:
                        if val.input_box_id == box.id:
                            value = val.value
                            break
                
                # Create a unique placeholder that won't be affected by markdown processing
                unique_placeholder = f'INPUTBOX_{box.id}_{box.box_id}'
                box_mappings[unique_placeholder] = {
                    'id': box.id,
                    'label': box.label,
                    'value': value
                }
                
                # Replace the original placeholder with our unique one
                processed_content = processed_content.replace(placeholder, unique_placeholder)
        
        # Use more extensions for better markdown rendering
        extensions = [
            'markdown.extensions.extra',
            'markdown.extensions.codehilite',
            'markdown.extensions.nl2br',  # Convert line breaks to <br>
            'markdown.extensions.sane_lists',
            'markdown.extensions.smarty',
            'markdown.extensions.tables',
        ]
        
        html = markdown.markdown(processed_content, extensions=extensions)
        
        # Use bleach to clean but allow most HTML tags for proper rendering
        # Convert frozenset to list before adding additional tags
        additional_tags = [
            'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'span', 'br', 'hr',
            'table', 'thead', 'tbody', 'tr', 'th', 'td', 'pre', 'code', 'blockquote'
        ]
        allowed_tags = list(bleach.sanitizer.ALLOWED_TAGS) + additional_tags
        
        allowed_attributes = {
            **bleach.sanitizer.ALLOWED_ATTRIBUTES,
            'img': ['src', 'alt', 'title'],
            'span': ['class', 'id', 'style', 'data-input-id', 'data-box-id', 'data-box-label', 'data-value', 
                     'contenteditable', 'title', 'onclick'],
            '*': ['class', 'id', 'style'],
        }
        
        clean_html = bleach.clean(
            html,
            tags=allowed_tags,
            attributes=allowed_attributes,
            strip=True
        )
        
        # Replace our unique placeholders with interactive spans
        if box_mappings:
            for placeholder, box in box_mappings.items():
                # Create a more visually distinct and clickable input box span
                input_span = f'<span class="interactive-input-box" ' + \
                           f'data-box-id="{box["id"]}" ' + \
                           f'data-box-label="{box["label"]}" ' + \
                           f'data-value="{box["value"]}" ' + \
                           f'title="Click to edit: {box["label"]}">' + \
                           f'{box["value"] or "[Click to enter value]"}</span>'
                clean_html = clean_html.replace(placeholder, input_span)
        
        return clean_html
    except Exception as e:
        # If there's any error during markdown rendering, return a simple error message
        error_html = f"<div class='alert alert-danger'>Error rendering content: {str(e)}</div>"
        return error_html

def create_docx(document_id):
    """
    Create a docx file from a document with improved markdown formatting
    """
    db_session = get_session(engine)
    document = db_session.query(Document).get(document_id)
    
    if not document:
        db_session.close()
        return None
    
    # Get rendered content
    content = document.get_rendered_content()
    
    # Create docx
    docx = DocxDocument()
    docx.add_heading(document.title, 0)
    
    # Process markdown content with better formatting
    lines = content.split('\n')
    current_paragraph = []
    in_list = False
    list_items = []
    
    for line in lines:
        # Handle headers
        if line.startswith('#'):
            # Add any accumulated paragraph text
            if current_paragraph:
                p = docx.add_paragraph(''.join(current_paragraph))
                current_paragraph = []
            
            # Add header with appropriate level
            level = len(re.match(r'^#+', line).group(0))
            header_text = line.lstrip('#').strip()
            docx.add_heading(header_text, level)
            continue
        
        # Handle list items
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            if not in_list:
                # Add any accumulated paragraph text before starting list
                if current_paragraph:
                    p = docx.add_paragraph(''.join(current_paragraph))
                    current_paragraph = []
                in_list = True
            
            list_items.append(line.strip()[2:].strip())
            continue
        elif in_list and line.strip() == '':
            # End of list, add the list items
            for item in list_items:
                p = docx.add_paragraph(item, style='List Bullet')
            list_items = []
            in_list = False
            continue
        elif in_list:
            # Continuation of a list item (indented content)
            if line.strip():
                list_items[-1] += ' ' + line.strip()
            continue
        
        # Handle regular paragraph text
        if line.strip() == '' and current_paragraph:
            # Empty line marks the end of a paragraph
            p = docx.add_paragraph(''.join(current_paragraph))
            current_paragraph = []
        elif line.strip():
            # Process inline markdown in the line
            formatted_line = line
            # Handle bold text
            formatted_line = re.sub(r'\*\*(.*?)\*\*', lambda m: m.group(1), formatted_line)
            # Handle italic text
            formatted_line = re.sub(r'\*(.*?)\*', lambda m: m.group(1), formatted_line)
            # Handle links
            formatted_line = re.sub(r'\[(.*?)\]\((.*?)\)', lambda m: m.group(1), formatted_line)
            
            current_paragraph.append(formatted_line + ' ')
    
    # Add any remaining paragraph text
    if current_paragraph:
        p = docx.add_paragraph(''.join(current_paragraph))
    
    # Add any remaining list items
    if list_items:
        for item in list_items:
            p = docx.add_paragraph(item, style='List Bullet')
    
    db_session.close()
    return docx

# Routes
@app.route('/')
def index():
    """Home page"""
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Login page"""
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        if not username or not password:
            flash('Please provide both username and password', 'danger')
            return render_template('login.html')
        
        db_session = get_session(engine)
        user = db_session.query(User).filter_by(username=username).first()
        
        if user and pbkdf2_sha256.verify(password, user.password_hash):
            login_user(user)
            db_session.close()
            return redirect(url_for('dashboard'))
        
        flash('Invalid username or password', 'danger')
        db_session.close()
    
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    """Registration page"""
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')
        
        if not username or not email or not password or not confirm_password:
            flash('All fields are required', 'danger')
            return render_template('register.html')
        
        if password != confirm_password:
            flash('Passwords do not match', 'danger')
            return render_template('register.html')
        
        db_session = get_session(engine)
        
        # Check if username or email already exists
        existing_user = db_session.query(User).filter(
            (User.username == username) | (User.email == email)
        ).first()
        
        if existing_user:
            flash('Username or email already exists', 'danger')
            db_session.close()
            return render_template('register.html')
        
        # Create new user
        new_user = User(
            username=username,
            email=email,
            password_hash=pbkdf2_sha256.hash(password),
            is_admin=False,
            is_active=True
        )
        
        db_session.add(new_user)
        db_session.commit()
        db_session.close()
        
        flash('Registration successful. You can now log in.', 'success')
        return redirect(url_for('login'))
    
    return render_template('register.html')

@app.route('/logout')
@login_required
def logout():
    """Logout route"""
    logout_user()
    return redirect(url_for('index'))

@app.route('/admin/users')
@login_required
@admin_required
def admin_users():
    """Admin user management page"""
    db_session = get_session(engine)
    
    # Get all users with their station assignments
    users = db_session.query(User).all()
    
    # Get all stations for the dropdown
    stations = db_session.query(Station).all()
    
    # Create a dictionary to store user's stations
    user_stations = {}
    
    # For each user, get their station assignments
    for user in users:
        # Query to get stations and roles for this user
        stations_data = db_session.query(Station, StationUser.role).join(
            StationUser, Station.id == StationUser.station_id
        ).filter(
            StationUser.user_id == user.id
        ).all()
        
        # Store as a list of dictionaries for easier template handling
        user_stations[user.id] = [
            {
                'id': station.id,
                'name': station.name,
                'role': role
            }
            for station, role in stations_data
        ]
    
    db_session.close()
    
    return render_template('admin_users.html', users=users, stations=stations, user_stations=user_stations)

@app.route('/admin/users/<int:user_id>/suspend')
@login_required
@admin_required
def admin_suspend_user(user_id):
    """Suspend a user"""
    db_session = get_session(engine)
    user = db_session.query(User).get(user_id)
    
    if not user:
        db_session.close()
        flash('User not found', 'danger')
        return redirect(url_for('admin_users'))
    
    # Prevent suspending the admin account
    if user.username == 'admin':
        db_session.close()
        flash('Cannot suspend admin account', 'danger')
        return redirect(url_for('admin_users'))
    
    user.is_active = False
    db_session.commit()
    db_session.close()
    
    flash(f'User {user.username} has been suspended', 'success')
    return redirect(url_for('admin_users'))

@app.route('/admin/users/<int:user_id>/activate')
@login_required
@admin_required
def admin_activate_user(user_id):
    """Activate a suspended user"""
    db_session = get_session(engine)
    user = db_session.query(User).get(user_id)
    
    if not user:
        db_session.close()
        flash('User not found', 'danger')
        return redirect(url_for('admin_users'))
    
    user.is_active = True
    db_session.commit()
    db_session.close()
    
    flash(f'User {user.username} has been activated', 'success')
    return redirect(url_for('admin_users'))

@app.route('/admin/users/<int:user_id>/delete')
@login_required
@admin_required
def admin_delete_user(user_id):
    """Delete a user"""
    db_session = get_session(engine)
    user = db_session.query(User).get(user_id)
    
    if not user:
        db_session.close()
        flash('User not found', 'danger')
        return redirect(url_for('admin_users'))
    
    # Prevent deleting the admin account
    if user.username == 'admin':
        db_session.close()
        flash('Cannot delete admin account', 'danger')
        return redirect(url_for('admin_users'))
    
    # Delete all user's documents and templates
    db_session.query(Document).filter_by(creator_id=user.id).delete()
    db_session.query(Template).filter_by(creator_id=user.id).delete()
    
    # Delete the user
    db_session.delete(user)
    db_session.commit()
    db_session.close()
    
    flash(f'User {user.username} has been deleted', 'success')
    return redirect(url_for('admin_users'))

@app.route('/dashboard')
@login_required
def dashboard():
    """User dashboard showing stations, templates and documents"""
    db_session = get_session(engine)
    
    # Get user's stations with their roles
    # Include stations where the user is a direct member
    user_stations = db_session.query(Station, StationUser.role).join(
        StationUser, Station.id == StationUser.station_id
    ).filter(
        StationUser.user_id == current_user.id
    ).all()

    # Convert to a list of dictionaries for easier template handling
    stations = []
    for station, role in user_stations:
        stations.append({
            'id': station.id,
            'name': station.name,
            'description': station.description,
            'role': role,
            'created_at': station.created_at
        })

    # Include all stations for admin users
    if current_user.is_admin:
        # Get any stations that the admin isn't directly a member of
        admin_stations = db_session.query(Station).filter(
            ~Station.id.in_([s['id'] for s in stations])
        ).all()
        
        # Add them to the list with the role 'Admin'
        for station in admin_stations:
            stations.append({
                'id': station.id,
                'name': station.name,
                'description': station.description,
                'role': 'Admin',
                'created_at': station.created_at
            })
    
    # Get user's templates
    my_templates = db_session.query(Template).filter_by(creator_id=current_user.id).all()
    
    # Get templates assigned to the user with eager loading of creator and assignments relationships
    assigned_templates_query = db_session.query(Template).options(
        joinedload(Template.creator),
        joinedload(Template.assignments)
    ).join(
        TemplateAssignment, Template.id == TemplateAssignment.template_id
    ).filter(
        TemplateAssignment.assignee_id == current_user.id
    )
    assigned_templates = assigned_templates_query.all()
    
    # Get station templates based on user's stations
    station_templates = db_session.query(Template).options(
        joinedload(Template.creator),
        joinedload(Template.station)
    ).filter(
        Template.station_id.in_([station['id'] for station in stations])
    ).all()
    
    # Get user's documents, eagerly loading the template relationship to prevent DetachedInstanceError
    documents = db_session.query(Document).options(joinedload(Document.template)).filter_by(creator_id=current_user.id).all()
    
    # Get all users for template assignment
    users = db_session.query(User).filter(User.id != current_user.id).all()
    
    db_session.close()
    
    return render_template('dashboard.html', 
                          stations=stations,
                          my_templates=my_templates, 
                          assigned_templates=assigned_templates,
                          station_templates=station_templates,
                          documents=documents,
                          users=users)

# Station Management Routes
@app.route('/admin/stations')
@login_required
@admin_required
def admin_stations():
    """Admin station management page"""
    db_session = get_session(engine)
    
    # Get all stations with their creators
    stations = db_session.query(Station).options(joinedload(Station.creator)).all()
    
    # Get all users for station assignment
    users = db_session.query(User).filter(User.is_active == True).all()
    
    db_session.close()
    
    return render_template('admin_stations.html', stations=stations, users=users)

@app.route('/admin/stations/create', methods=['GET', 'POST'])
@login_required
@admin_required
def create_station():
    """Create a new station"""
    if request.method == 'POST':
        name = request.form.get('name')
        description = request.form.get('description')
        
        if not name:
            flash('Station name is required', 'danger')
            return redirect(url_for('admin_stations'))
        
        db_session = get_session(engine)
        
        # Check if station name already exists
        existing_station = db_session.query(Station).filter_by(name=name).first()
        if existing_station:
            db_session.close()
            flash('A station with this name already exists', 'danger')
            return redirect(url_for('admin_stations'))
        
        # Create station
        new_station = Station(
            name=name,
            description=description,
            created_by=current_user.id
        )
        
        db_session.add(new_station)
        db_session.commit()
        
        # Get station ID for redirection
        station_id = new_station.id
        
        db_session.close()
        
        flash('Station created successfully', 'success')
        return redirect(url_for('admin_stations'))
    
    return render_template('create_station.html')

@app.route('/admin/stations/<int:station_id>/delete')
@login_required
@admin_required
def delete_station(station_id):
    """Delete a station"""
    db_session = get_session(engine)
    station = db_session.query(Station).get(station_id)
    
    if not station:
        db_session.close()
        flash('Station not found', 'danger')
        return redirect(url_for('admin_stations'))
    
    # Check if there are templates in this station
    templates_count = db_session.query(Template).filter_by(station_id=station_id).count()
    
    if templates_count > 0:
        db_session.close()
        flash(f'Cannot delete station: {templates_count} templates are assigned to it', 'danger')
        return redirect(url_for('admin_stations'))
    
    # Delete all user assignments to this station
    db_session.query(StationUser).filter_by(station_id=station_id).delete()
    
    # Delete the station
    db_session.delete(station)
    db_session.commit()
    db_session.close()
    
    flash('Station deleted successfully', 'success')
    return redirect(url_for('admin_stations'))

@app.route('/admin/stations/<int:station_id>/edit', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_station(station_id):
    """Edit a station"""
    db_session = get_session(engine)
    station = db_session.query(Station).get(station_id)
    
    if not station:
        db_session.close()
        flash('Station not found', 'danger')
        return redirect(url_for('admin_stations'))
    
    if request.method == 'POST':
        name = request.form.get('name')
        description = request.form.get('description')
        
        if not name:
            flash('Station name is required', 'danger')
            return render_template('edit_station.html', station=station)
        
        # Check if the new name conflicts with another station
        if name != station.name:
            existing_station = db_session.query(Station).filter_by(name=name).first()
            if existing_station:
                flash('A station with this name already exists', 'danger')
                return render_template('edit_station.html', station=station)
        
        # Update station
        station.name = name
        station.description = description
        
        db_session.commit()
        db_session.close()
        
        flash('Station updated successfully', 'success')
        return redirect(url_for('admin_stations'))
    
    db_session.close()
    return render_template('edit_station.html', station=station)

@app.route('/admin/stations/<int:station_id>/users')
@login_required
@admin_required
def station_users(station_id):
    """Manage users in a station"""
    db_session = get_session(engine)
    station = db_session.query(Station).get(station_id)
    
    if not station:
        db_session.close()
        flash('Station not found', 'danger')
        return redirect(url_for('admin_stations'))
    
    # Get all users in this station with their roles
    station_users = db_session.query(User, StationUser.role).join(
        StationUser, User.id == StationUser.user_id
    ).filter(
        StationUser.station_id == station_id
    ).all()
    
    # Get all active users that are not already in the station
    # Convert the list of users in the station to a list of user IDs
    station_user_ids = [user.id for user, _ in station_users]
    
    # Get all active users except those already in the station
    # Fixed bug: Using User.id.notin_() instead of ~User.id.in_() for empty list case
    if station_user_ids:
        available_users = db_session.query(User).filter(
            User.is_active == True,
            User.id.notin_(station_user_ids)
        ).all()
    else:
        available_users = db_session.query(User).filter(
            User.is_active == True
        ).all()
    
    print(f"Station users: {len(station_users)}")
    print(f"Available users: {len(available_users)}")
    
    # Store station name before closing session
    station_name = station.name
    
    db_session.close()
    
    return render_template('station_users.html', 
                          station=station, 
                          station_users=station_users, 
                          available_users=available_users,
                          roles=[StationUserRole.STATION_MASTER.value, StationUserRole.STATION_STAFF.value])

@app.route('/admin/stations/<int:station_id>/users/add', methods=['POST'])
@login_required
@admin_required
def add_station_user(station_id):
    """Add a user to a station"""
    db_session = get_session(engine)
    station = db_session.query(Station).get(station_id)
    
    if not station:
        db_session.close()
        flash('Station not found', 'danger')
        return redirect(url_for('admin_stations'))
    
    user_id = request.form.get('user_id')
    role = request.form.get('role')
    
    if not user_id or not role:
        db_session.close()
        flash('User and role are required', 'danger')
        return redirect(url_for('station_users', station_id=station_id))
    
    # Check if the user exists
    user = db_session.query(User).get(user_id)
    if not user:
        db_session.close()
        flash('User not found', 'danger')
        return redirect(url_for('station_users', station_id=station_id))
    
    # Check if user is already in the station with the same role
    existing = db_session.query(StationUser).filter_by(
        station_id=station_id,
        user_id=user_id,
        role=role
    ).first()
    
    if existing:
        db_session.close()
        flash(f'User is already in this station with role {role}', 'danger')
        return redirect(url_for('station_users', station_id=station_id))
    
    # Check if user is in the station with a different role
    existing_different_role = db_session.query(StationUser).filter_by(
        station_id=station_id,
        user_id=user_id
    ).first()
    
    if existing_different_role:
        # Update the role instead of creating a new record
        existing_different_role.role = role
        db_session.commit()
        flash(f'User {user.username} role updated to {role}', 'success')
        
        # Create notification for the user
        notification = Notification(
            user_id=user_id,
            title="Role Changed in Station",
            message=f"Your role in the station '{station.name}' has been changed to {role}.",
            notification_type="station_role_change"
        )
        
        db_session.add(notification)
        db_session.commit()
        db_session.close()
        
        return redirect(url_for('station_users', station_id=station_id))
    
    # Add user to station
    station_user = StationUser(
        station_id=station_id,
        user_id=user_id,
        role=role
    )
    
    db_session.add(station_user)
    db_session.commit()
    
    # Create notification for the user
    notification = Notification(
        user_id=user_id,
        title="Added to Station",
        message=f"You have been added to the station '{station.name}' with the role of {role}.",
        notification_type="station_assignment"
    )
    
    db_session.add(notification)
    db_session.commit()
    db_session.close()
    
    flash(f'User {user.username} added to station with role {role}', 'success')
    return redirect(url_for('station_users', station_id=station_id))

@app.route('/admin/stations/<int:station_id>/users/<int:user_id>/remove')
@login_required
@admin_required
def remove_station_user(station_id, user_id):
    """Remove a user from a station"""
    db_session = get_session(engine)
    
    # Check if the station and user exist
    station = db_session.query(Station).get(station_id)
    user = db_session.query(User).get(user_id)
    
    if not station or not user:
        db_session.close()
        flash('Station or user not found', 'danger')
        return redirect(url_for('admin_stations'))
    
    # Remove the user from the station
    station_user = db_session.query(StationUser).filter_by(
        station_id=station_id,
        user_id=user_id
    ).first()
    
    if station_user:
        db_session.delete(station_user)
        
        # Create notification for the user
        notification = Notification(
            user_id=user_id,
            title="Removed from Station",
            message=f"You have been removed from the station '{station.name}'.",
            notification_type="station_removal"
        )
        
        db_session.add(notification)
        db_session.commit()
        
        flash(f'User {user.username} removed from station', 'success')
    else:
        flash('User is not in this station', 'danger')
    
    db_session.close()
    return redirect(url_for('station_users', station_id=station_id))

@app.route('/admin/stations/<int:station_id>/users/<int:user_id>/change-role', methods=['POST'])
@login_required
@admin_required
def change_station_user_role(station_id, user_id):
    """Change a user's role in a station"""
    db_session = get_session(engine)
    
    # Check if the station and user exist
    station = db_session.query(Station).get(station_id)
    user = db_session.query(User).get(user_id)
    
    if not station or not user:
        db_session.close()
        flash('Station or user not found', 'danger')
        return redirect(url_for('admin_stations'))
    
    role = request.form.get('role')
    if not role:
        db_session.close()
        flash('Role is required', 'danger')
        return redirect(url_for('station_users', station_id=station_id))
    
    # Update the user's role
    station_user = db_session.query(StationUser).filter_by(
        station_id=station_id,
        user_id=user_id
    ).first()
    
    if station_user:
        station_user.role = role
        
        # Create notification for the user
        notification = Notification(
            user_id=user_id,
            title="Role Changed in Station",
            message=f"Your role in the station '{station.name}' has been changed to {role}.",
            notification_type="station_role_change"
        )
        
        db_session.add(notification)
        db_session.commit()
        
        flash(f'User {user.username} role changed to {role}', 'success')
    else:
        flash('User is not in this station', 'danger')
    
    db_session.close()
    return redirect(url_for('station_users', station_id=station_id))

@app.route('/admin/assign-multiple-users', methods=['POST'])
@login_required
@admin_required
def assign_multiple_users_to_station():
    """Assign multiple users to a station at once"""
    db_session = get_session(engine)
    
    station_id = request.form.get('station_id')
    role = request.form.get('role')
    user_ids = request.form.get('user_ids')
    
    if not station_id or not role or not user_ids:
        db_session.close()
        flash('Station, role and users are required', 'danger')
        return redirect(url_for('admin_users'))
    
    # Check if the station exists
    station = db_session.query(Station).get(station_id)
    if not station:
        db_session.close()
        flash('Station not found', 'danger')
        return redirect(url_for('admin_users'))
    
    # Store station name before session closes
    station_name = station.name
    
    # Split user_ids by comma
    user_id_list = user_ids.split(',')
    assigned_count = 0
    updated_count = 0
    
    for user_id in user_id_list:
        try:
            user_id = int(user_id.strip())
            
            # Check if user exists
            user = db_session.query(User).get(user_id)
            if not user:
                continue
            
            # Check if user is already in the station
            existing_user = db_session.query(StationUser).filter_by(
                station_id=station_id,
                user_id=user_id
            ).first()
            
            if existing_user:
                # Update role if it's different
                if existing_user.role != role:
                    existing_user.role = role
                    updated_count += 1
                    
                    # Create notification for role change
                    notification = Notification(
                        user_id=user_id,
                        title="Role Changed in Station",
                        message=f"Your role in the station '{station_name}' has been changed to {role}.",
                        notification_type="station_role_change"
                    )
                    db_session.add(notification)
            else:
                # Add user to station
                station_user = StationUser(
                    station_id=station_id,
                    user_id=user_id,
                    role=role
                )
                db_session.add(station_user)
                assigned_count += 1
                
                # Create notification for new assignment
                notification = Notification(
                    user_id=user_id,
                    title="Added to Station",
                    message=f"You have been added to the station '{station_name}' with the role of {role}.",
                    notification_type="station_assignment"
                )
                db_session.add(notification)
        except Exception as e:
            print(f"Error processing user {user_id}: {str(e)}")
            continue
    
    db_session.commit()
    db_session.close()
    
    if assigned_count > 0 and updated_count > 0:
        flash(f'{assigned_count} users added and {updated_count} users updated in station {station_name}', 'success')
    elif assigned_count > 0:
        flash(f'{assigned_count} users added to station {station_name}', 'success')
    elif updated_count > 0:
        flash(f'{updated_count} users updated in station {station_name}', 'success')
    else:
        flash('No users were added or updated', 'warning')
    
    return redirect(url_for('admin_users'))

@app.route('/stations/<int:station_id>')
@login_required
def view_station(station_id):
    """View a station and its templates"""
    db_session = get_session(engine)
    
    # Check if the station exists
    station = db_session.query(Station).get(station_id)
    
    if not station:
        db_session.close()
        flash('Station not found', 'danger')
        return redirect(url_for('dashboard'))
    
    # Check if the user is in this station
    station_user = db_session.query(StationUser).filter_by(
        station_id=station_id,
        user_id=current_user.id
    ).first()
    
    # Allow access for admins even if they're not in the station
    if not station_user and not current_user.is_admin:
        db_session.close()
        flash('You do not have access to this station', 'danger')
        return redirect(url_for('dashboard'))
    
    # Get the user's role (for admins who might not be in the station)
    user_role = station_user.role if station_user else "Admin"
    
    # Get templates in this station
    templates = db_session.query(Template).options(
        joinedload(Template.creator)
    ).filter(
        Template.station_id == station_id
    ).all()
    
    # Get all users in this station with their roles
    station_users = db_session.query(User, StationUser.role).join(
        StationUser, User.id == StationUser.user_id
    ).filter(
        StationUser.station_id == station_id
    ).all()
    
    db_session.close()
    
    return render_template('view_station.html', 
                          station=station, 
                          user_role=user_role,
                          templates=templates,
                          station_users=station_users)

@app.route('/templates/create', methods=['GET', 'POST'])
@login_required
def create_template():
    """Create a new template"""
    db_session = get_session(engine)
    
    # Check if a station_id is provided in the query parameters
    # This indicates creating a template from a specific station view
    from_station_id = request.args.get('station_id')
    
    if from_station_id and not current_user.is_admin:
        # If creating from a specific station, only allow that station to be selected
        station = db_session.query(Station).get(from_station_id)
        if station:
            # Verify user is station master for this station
            station_user = db_session.query(StationUser).filter_by(
                station_id=from_station_id,
                user_id=current_user.id,
                role=StationUserRole.STATION_MASTER.value
            ).first()
            
            if station_user:
                stations = [{
                    'id': station.id,
                    'name': station.name
                }]
            else:
                # User is not a station master for this station
                stations = []
        else:
            stations = []
    else:
        # Get all stations the user has access to for template creation
        user_stations = db_session.query(Station, StationUser.role).join(
            StationUser, Station.id == StationUser.station_id
        ).filter(
            StationUser.user_id == current_user.id
        ).all()
        
        # Convert to list of dictionaries and filter for stations where user is master
        stations = []
        for station, role in user_stations:
            # Station masters and admins can create templates in a station
            if role == StationUserRole.STATION_MASTER.value or current_user.is_admin:
                stations.append({
                    'id': station.id,
                    'name': station.name
                })
    
    db_session.close()
    
    if request.method == 'POST':
        title = request.form.get('title')
        description = request.form.get('description')
        content = request.form.get('content')
        station_id = request.form.get('station_id')
        
        if not title or not content:
            flash('Title and content are required', 'danger')
            return render_template('create_template.html', stations=stations)
        
        # Extract input boxes
        input_boxes = extract_input_boxes(content)
        
        db_session = get_session(engine)
        
        # If station_id is provided, validate it
        if station_id:
            # Check if station exists
            station = db_session.query(Station).get(station_id)
            if not station:
                db_session.close()
                flash('Selected station not found', 'danger')
                return render_template('create_template.html', stations=stations)
            
            # Check if user has permission to create templates in this station
            if not current_user.is_admin:
                station_user = db_session.query(StationUser).filter_by(
                    station_id=station_id,
                    user_id=current_user.id,
                    role=StationUserRole.STATION_MASTER.value
                ).first()
                
                if not station_user:
                    db_session.close()
                    flash('You do not have permission to create templates in this station', 'danger')
                    return render_template('create_template.html', stations=stations)
        
        # Create template
        new_template = Template(
            title=title,
            description=description,
            content=content,
            creator_id=current_user.id,
            station_id=station_id if station_id else None
        )
        
        db_session.add(new_template)
        db_session.flush()  # Flush to get the ID
        
        # Create input boxes
        for i, box_id in enumerate(input_boxes):
            input_box = InputBox(
                box_id=box_id,
                template_id=new_template.id,
                label=box_id,  # Use box_id as label by default
                position=i
            )
            db_session.add(input_box)
        
        db_session.commit()
        db_session.close()
        
        flash('Template created successfully', 'success')
        
        # Redirect to station view if template was created for a station
        if station_id:
            return redirect(url_for('view_station', station_id=station_id))
        else:
            return redirect(url_for('dashboard'))
    
    return render_template('create_template.html', stations=stations)

@app.route('/templates/<int:template_id>')
@login_required
def view_template(template_id):
    """View a template"""
    db_session = get_session(engine)
    # Eagerly load the input_boxes relationship to prevent DetachedInstanceError
    template = db_session.query(Template).options(joinedload(Template.input_boxes)).get(template_id)
    
    if not template:
        db_session.close()
        flash('Template not found', 'danger')
        return redirect(url_for('dashboard'))
    
    # Check if the user is the creator
    if template.creator_id != current_user.id:
        db_session.close()
        flash('You do not have permission to view this template', 'danger')
        return redirect(url_for('dashboard'))
    
    # Convert markdown to HTML
    html_content = markdown_to_html(template.content)
    
    db_session.close()
    
    return render_template('view_template.html', template=template, html_content=html_content)

@app.route('/templates/<int:template_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_template(template_id):
    """Edit a template"""
    db_session = get_session(engine)
    # Eagerly load the input_boxes relationship to prevent DetachedInstanceError
    template = db_session.query(Template).options(joinedload(Template.input_boxes)).get(template_id)
    
    if not template:
        db_session.close()
        flash('Template not found', 'danger')
        return redirect(url_for('dashboard'))
    
    # Check if the user is the creator or has edit permission
    has_permission = False
    
    # Creator always has permission
    if template.creator_id == current_user.id:
        has_permission = True
    else:
        # Check if user has been assigned this template with edit permission
        assignment = db_session.query(TemplateAssignment).filter_by(
            template_id=template.id,
            assignee_id=current_user.id,
            can_edit=True
        ).first()
        
        if assignment:
            has_permission = True
    
    if not has_permission:
        db_session.close()
        flash('You do not have permission to edit this template', 'danger')
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        title = request.form.get('title')
        description = request.form.get('description')
        content = request.form.get('content')
        
        if not title or not content:
            flash('Title and content are required', 'danger')
            return render_template('edit_template.html', template=template)
        
        # Extract input boxes
        new_input_boxes = extract_input_boxes(content)
        
        # Update template
        template.title = title
        template.description = description
        template.content = content
        
        # Delete old input boxes and create new ones
        db_session.query(InputBox).filter_by(template_id=template.id).delete()
        db_session.flush()
        
        # Create input boxes
        for i, box_id in enumerate(new_input_boxes):
            input_box = InputBox(
                box_id=box_id,
                template_id=template.id,
                label=box_id,  # Use box_id as label by default
                position=i
            )
            db_session.add(input_box)
        
        db_session.commit()
        
        flash('Template updated successfully', 'success')
        return redirect(url_for('view_template', template_id=template.id))
    
    db_session.close()
    
    return render_template('edit_template.html', template=template)

@app.route('/templates/<int:template_id>/create_document', methods=['GET', 'POST'])
@login_required
def create_document(template_id):
    """Create a document from a template"""
    db_session = get_session(engine)
    # Eagerly load the input_boxes relationship to prevent DetachedInstanceError
    template = db_session.query(Template).options(joinedload(Template.input_boxes)).get(template_id)
    
    if not template:
        db_session.close()
        flash('Template not found', 'danger')
        return redirect(url_for('dashboard'))
    
    # Get input boxes for the template
    input_boxes = db_session.query(InputBox).filter_by(template_id=template.id).order_by(InputBox.position).all()
    
    if request.method == 'POST':
        title = request.form.get('title')
        description = request.form.get('description')
        
        if not title:
            flash('Title is required', 'danger')
            return render_template('create_document.html', template=template, input_boxes=input_boxes)
        
        # Create document
        document = Document(
            title=title,
            description=description,
            template_id=template.id,
            creator_id=current_user.id
        )
        
        db_session.add(document)
        db_session.flush()  # Flush to get the ID
        
        # Create input values
        for input_box in input_boxes:
            value = request.form.get(f'input_{input_box.id}', '')
            
            input_value = DocumentInputValue(
                document_id=document.id,
                input_box_id=input_box.id,
                value=value
            )
            db_session.add(input_value)
        
        db_session.commit()
        
        flash('Document created successfully', 'success')
        return redirect(url_for('view_document', document_id=document.id))
    
    db_session.close()
    
    return render_template('create_document.html', template=template, input_boxes=input_boxes)

@app.route('/documents/<int:document_id>')
@login_required
def view_document(document_id):
    """View a document"""
    db_session = get_session(engine)
    
    # Eagerly load the template relationship and input_values with input_box to prevent DetachedInstanceError
    document = db_session.query(Document).options(
        joinedload(Document.template),
        joinedload(Document.input_values).joinedload(DocumentInputValue.input_box)
    ).get(document_id)
    
    if not document:
        db_session.close()
        flash('Document not found', 'danger')
        return redirect(url_for('dashboard'))
    
    # Check if the user is the creator
    if document.creator_id != current_user.id:
        db_session.close()
        flash('You do not have permission to view this document', 'danger')
        return redirect(url_for('dashboard'))
    
    try:
        # Get input boxes for the template
        input_boxes = db_session.query(InputBox).filter_by(template_id=document.template_id).order_by(InputBox.position).all()
        
        # Get input values
        input_values = db_session.query(DocumentInputValue).filter_by(document_id=document.id).all()
        
        # Get the template content and rendered content
        template_content = document.template.content
        rendered_content = document.get_rendered_content()
        
        # Convert markdown to HTML with interactive input boxes
        # We use rendered_content for viewing since it already has the values filled in
        html_content = markdown_to_html(rendered_content, input_boxes, input_values, document_id)
    except Exception as e:
        db_session.close()
        flash(f'Error rendering document: {str(e)}', 'danger')
        return redirect(url_for('dashboard'))
    
    # Get input values and filter out any with null input_box
    input_values = [v for v in document.input_values if v.input_box is not None]
    
    # Create a dictionary with input_box_id -> value
    input_value_dict = {value.input_box_id: value for value in input_values}
    
    # Get input boxes
    input_boxes = db_session.query(InputBox).filter_by(template_id=document.template_id).order_by(InputBox.position).all()
    
    db_session.close()
    
    return render_template('view_document.html', 
                          document=document, 
                          html_content=html_content, 
                          input_boxes=input_boxes,
                          input_value_dict=input_value_dict)

@app.route('/documents/<int:document_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_document(document_id):
    """Edit a document"""
    db_session = get_session(engine)
    # Eagerly load the template relationship to prevent DetachedInstanceError
    document = db_session.query(Document).options(joinedload(Document.template)).get(document_id)
    
    if not document:
        db_session.close()
        flash('Document not found', 'danger')
        return redirect(url_for('dashboard'))
    
    # Check if the user is the creator
    if document.creator_id != current_user.id:
        db_session.close()
        flash('You do not have permission to edit this document', 'danger')
        return redirect(url_for('dashboard'))
    
    # Get input boxes for the template
    input_boxes = db_session.query(InputBox).filter_by(template_id=document.template_id).order_by(InputBox.position).all()
    
    # Get input values
    input_values = db_session.query(DocumentInputValue).filter_by(document_id=document.id).all()
    # Create a dictionary with input_box_id -> value
    input_value_dict = {value.input_box_id: value.value for value in input_values}
    
    if request.method == 'POST':
        title = request.form.get('title')
        description = request.form.get('description')
        
        if not title:
            flash('Title is required', 'danger')
            return render_template('edit_document.html', document=document, input_boxes=input_boxes, input_value_dict=input_value_dict)
        
        # Update document
        document.title = title
        document.description = description
        
        # Update input values
        for input_box in input_boxes:
            value = request.form.get(f'input_{input_box.id}', '')
            
            # Get existing value or create new one
            input_value = db_session.query(DocumentInputValue).filter_by(
                document_id=document.id,
                input_box_id=input_box.id
            ).first()
            
            if input_value:
                input_value.value = value
            else:
                input_value = DocumentInputValue(
                    document_id=document.id,
                    input_box_id=input_box.id,
                    value=value
                )
                db_session.add(input_value)
        
        db_session.commit()
        
        flash('Document updated successfully', 'success')
        return redirect(url_for('view_document', document_id=document.id))
    
    db_session.close()
    
    # For document editing, we need to ensure we're working with the template content
    # and replacing the input box placeholders with their current values
    template = db_session.query(Template).options(joinedload(Template.input_boxes)).get(document.template_id)
    template_content = template.content
    
    # Pass input boxes, values and document_id for interactive editing
    input_values = db_session.query(DocumentInputValue).filter_by(document_id=document.id).all()
    html_content = markdown_to_html(template_content, input_boxes, input_values, document.id)
    
    # Debug information for the console
    print(f"Rendering template for document {document.id} with {len(input_boxes)} input boxes")
    for box in input_boxes:
        value = input_value_dict.get(box.id, "")
        print(f"Input box {box.id} ({box.box_id}): {value}")
    
    return render_template('edit_document.html', 
                          document=document, 
                          input_boxes=input_boxes, 
                          input_value_dict=input_value_dict,
                          html_content=html_content)

@app.route('/documents/<int:document_id>/export/docx')
@login_required
def export_docx(document_id):
    """Export a document to docx"""
    db_session = get_session(engine)
    document = db_session.query(Document).get(document_id)
    
    if not document:
        db_session.close()
        flash('Document not found', 'danger')
        return redirect(url_for('dashboard'))
    
    # Check if the user is the creator
    if document.creator_id != current_user.id:
        db_session.close()
        flash('You do not have permission to export this document', 'danger')
        return redirect(url_for('dashboard'))
    
    # Create docx
    docx = create_docx(document_id)
    
    if not docx:
        db_session.close()
        flash('Error creating docx file', 'danger')
        return redirect(url_for('view_document', document_id=document_id))
    
    # Create export directory if it doesn't exist
    export_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "exports")
    if not os.path.exists(export_dir):
        os.makedirs(export_dir)
    
    # Save docx
    docx_path = os.path.join(export_dir, f"{document.title.replace(' ', '_')}.docx")
    docx.save(docx_path)
    
    db_session.close()
    
    flash(f'Document exported to {docx_path}', 'success')
    return redirect(url_for('view_document', document_id=document_id))

@app.route('/documents/<int:document_id>/export/txt')
@login_required
def export_txt(document_id):
    """Export a document to txt with improved formatting"""
    db_session = get_session(engine)
    document = db_session.query(Document).get(document_id)
    
    if not document:
        db_session.close()
        flash('Document not found', 'danger')
        return redirect(url_for('dashboard'))
    
    # Check if the user is the creator
    if document.creator_id != current_user.id:
        db_session.close()
        flash('You do not have permission to export this document', 'danger')
        return redirect(url_for('dashboard'))
    
    # Get rendered content
    content = document.get_rendered_content()
    
    # Process markdown for better TXT formatting
    lines = content.split('\n')
    processed_lines = []
    
    for line in lines:
        # Handle headers - keep them with some emphasis
        header_match = re.match(r'^(#+)\s+(.*?)$', line)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2)
            
            if level == 1:
                processed_lines.append('\n' + text.upper() + '\n' + '=' * len(text))
            elif level == 2:
                processed_lines.append('\n' + text + '\n' + '-' * len(text))
            else:
                processed_lines.append('\n' + text)
            continue
            
        # Process inline formatting
        processed_line = line
        processed_line = re.sub(r'\*\*(.*?)\*\*', r'\1', processed_line)  # Remove bold
        processed_line = re.sub(r'\*(.*?)\*', r'\1', processed_line)  # Remove italic
        processed_line = re.sub(r'!\[(.*?)\]\((.*?)\)', '', processed_line)  # Remove images
        processed_line = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1', processed_line)  # Remove links
        
        # Handle list items - keep the bullets
        if processed_line.strip().startswith('- ') or processed_line.strip().startswith('* '):
            # Preserve list formatting with proper indentation
            processed_lines.append(processed_line)
        else:
            processed_lines.append(processed_line)
    
    # Join lines back together
    plain_content = '\n'.join(processed_lines)
    
    # Create export directory if it doesn't exist
    export_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "exports")
    if not os.path.exists(export_dir):
        os.makedirs(export_dir)
    
    # Save txt
    txt_path = os.path.join(export_dir, f"{document.title.replace(' ', '_')}.txt")
    with open(txt_path, 'w') as f:
        f.write(plain_content)
    
    db_session.close()
    
    flash(f'Document exported to {txt_path}', 'success')
    return redirect(url_for('view_document', document_id=document_id))

@app.route('/templates/<int:template_id>/assign', methods=['POST'])
@login_required
def assign_template(template_id):
    """Assign a template to another user"""
    db_session = get_session(engine)
    template = db_session.query(Template).get(template_id)
    
    if not template:
        db_session.close()
        flash('Template not found', 'danger')
        return redirect(url_for('dashboard'))
    
    # Check if the user is the creator
    if template.creator_id != current_user.id:
        db_session.close()
        flash('You do not have permission to assign this template', 'danger')
        return redirect(url_for('dashboard'))
    
    # Get form data
    assignee_id = request.form.get('assignee_id')
    can_edit = 'can_edit' in request.form
    
    if not assignee_id:
        db_session.close()
        flash('Please select a user to assign the template to', 'danger')
        return redirect(url_for('dashboard'))
    
    # Check if the assignee exists
    assignee = db_session.query(User).get(assignee_id)
    if not assignee:
        db_session.close()
        flash('Selected user not found', 'danger')
        return redirect(url_for('dashboard'))
    
    # Check if assignment already exists
    existing_assignment = db_session.query(TemplateAssignment).filter_by(
        template_id=template.id,
        assignee_id=assignee.id
    ).first()
    
    if existing_assignment:
        # Update existing assignment
        existing_assignment.can_edit = can_edit
        existing_assignment.assigned_by_id = current_user.id
        existing_assignment.assigned_at = datetime.datetime.utcnow()
        
        # Save the username before closing the session
        username = assignee.username
        
        db_session.commit()
        db_session.close()
        
        flash(f'Template reassigned to {username}', 'success')
        return redirect(url_for('dashboard'))
    
    # Create new assignment
    assignment = TemplateAssignment(
        template_id=template.id,
        assignee_id=assignee.id,
        assigned_by_id=current_user.id,
        can_edit=can_edit
    )
    
    # Save the username before closing the session
    username = assignee.username
    
    db_session.add(assignment)
    db_session.commit()
    db_session.close()
    
    flash(f'Template assigned to {username}', 'success')
    return redirect(url_for('dashboard'))

@app.route('/templates/<int:template_id>/fork')
@login_required
def fork_template(template_id):
    """Fork a template to create your own copy"""
    db_session = get_session(engine)
    
    # Load the template with input boxes
    template = db_session.query(Template).options(joinedload(Template.input_boxes)).get(template_id)
    
    if not template:
        db_session.close()
        flash('Template not found', 'danger')
        return redirect(url_for('dashboard'))
    
    # Check if user has access to the template (creator, assignee, or station member)
    has_access = False
    
    # Case 1: User is the creator
    if template.creator_id == current_user.id:
        has_access = True
    
    # Case 2: User has been assigned this template
    if not has_access:
        assignment = db_session.query(TemplateAssignment).filter_by(
            template_id=template.id,
            assignee_id=current_user.id
        ).first()
        
        if assignment:
            has_access = True
    
    # Case 3: User is a member of the template's station (new feature)
    if not has_access and template.station_id:
        # Check if user is a member of this station
        station_user = db_session.query(StationUser).filter_by(
            station_id=template.station_id,
            user_id=current_user.id
        ).first()
        
        if station_user:
            has_access = True
    
    # Case 4: User is admin
    if not has_access and current_user.is_admin:
        has_access = True
    
    if not has_access:
        db_session.close()
        flash('You do not have permission to fork this template', 'danger')
        return redirect(url_for('dashboard'))
    
    # Create a new template as a fork
    forked_template = Template(
        title=f"{template.title} (Fork)",
        description=template.description,
        content=template.content,
        creator_id=current_user.id,
        parent_id=template.id,
        is_fork=True
    )
    
    db_session.add(forked_template)
    db_session.flush()  # Flush to get the ID
    
    # Copy input boxes
    for box in template.input_boxes:
        new_box = InputBox(
            box_id=box.box_id,
            template_id=forked_template.id,
            label=box.label,
            default_value=box.default_value,
            position=box.position
        )
        db_session.add(new_box)
    
    db_session.commit()
    
    # Store the template_id before closing the session to prevent DetachedInstanceError
    forked_template_id = forked_template.id
    
    db_session.close()
    
    flash('Template forked successfully', 'success')
    return redirect(url_for('view_template', template_id=forked_template_id))

@app.route('/api/input_value/<int:document_id>/<int:input_box_id>', methods=['POST'])
@login_required
def update_input_value(document_id, input_box_id):
    """API to update an input value"""
    db_session = get_session(engine)
    document = db_session.query(Document).get(document_id)
    
    if not document:
        db_session.close()
        return jsonify({'status': 'error', 'message': 'Document not found'}), 404
    
    # Check if the user is the creator
    if document.creator_id != current_user.id:
        db_session.close()
        return jsonify({'status': 'error', 'message': 'Permission denied'}), 403
    
    # Get input box
    input_box = db_session.query(InputBox).get(input_box_id)
    
    if not input_box or input_box.template_id != document.template_id:
        db_session.close()
        return jsonify({'status': 'error', 'message': 'Input box not found'}), 404
    
    # Get data from request
    data = request.get_json()
    value = data.get('value', '')
    
    # Sanitize input
    value = bleach.clean(value)
    
    # Get existing value or create new one
    input_value = db_session.query(DocumentInputValue).filter_by(
        document_id=document.id,
        input_box_id=input_box.id
    ).first()
    
    if input_value:
        input_value.value = value
    else:
        input_value = DocumentInputValue(
            document_id=document.id,
            input_box_id=input_box.id,
            value=value
        )
        db_session.add(input_value)
    
    db_session.commit()
    db_session.close()
    
    return jsonify({'status': 'success', 'message': 'Input value updated'})

@app.route('/templates/<int:template_id>/delete')
@login_required
def delete_template(template_id):
    """Delete a template and all associated documents"""
    db_session = get_session(engine)
    
    # Get the template with eager loading of input_boxes
    template = db_session.query(Template).options(joinedload(Template.input_boxes)).get(template_id)
    
    if not template:
        db_session.close()
        flash('Template not found', 'danger')
        return redirect(url_for('dashboard'))
    
    # Check if the user is the creator
    if template.creator_id != current_user.id:
        db_session.close()
        flash('You do not have permission to delete this template', 'danger')
        return redirect(url_for('dashboard'))
    
    try:
        # Get all documents based on this template
        documents = db_session.query(Document).filter_by(template_id=template.id).all()
        
        # Find all users who have documents based on this template
        affected_users = set()
        for doc in documents:
            if doc.creator_id != current_user.id:  # Don't notify yourself
                affected_users.add(doc.creator_id)
        
        # Get template name before deletion
        template_name = template.title
        
        # Create notifications for affected users
        for user_id in affected_users:
            notification = Notification(
                user_id=user_id,
                title="Template Deleted",
                message=f"A template you were using '{template_name}' has been deleted by {current_user.username}. Any documents you created using this template have also been deleted.",
                notification_type="template_deletion"
            )
            db_session.add(notification)
        
        # Delete all documents based on this template
        for document in documents:
            # Delete document input values first
            db_session.query(DocumentInputValue).filter_by(document_id=document.id).delete()
        
        # Now delete the documents
        db_session.query(Document).filter_by(template_id=template.id).delete()
        
        # Delete template assignments
        db_session.query(TemplateAssignment).filter_by(template_id=template.id).delete()
        
        # Delete the template
        db_session.delete(template)
        
        db_session.commit()
        
        flash(f'Template "{template_name}" and all associated documents have been deleted', 'success')
        
    except Exception as e:
        db_session.rollback()
        flash(f'Error deleting template: {str(e)}', 'danger')
    finally:
        db_session.close()
    
    return redirect(url_for('dashboard'))

@app.route('/notifications')
@login_required
def view_notifications():
    """View all notifications for the current user"""
    db_session = get_session(engine)
    
    # Get all notifications for the current user
    notifications = db_session.query(Notification).filter_by(
        user_id=current_user.id
    ).order_by(Notification.created_at.desc()).all()
    
    db_session.close()
    
    return render_template('notifications.html', notifications=notifications)

@app.route('/notifications/<int:notification_id>/mark-as-read')
@login_required
def mark_notification_read(notification_id):
    """Mark a notification as read"""
    db_session = get_session(engine)
    
    # Get the notification
    notification = db_session.query(Notification).filter_by(
        id=notification_id,
        user_id=current_user.id
    ).first()
    
    if notification:
        notification.is_read = True
        db_session.commit()
    
    db_session.close()
    
    return redirect(url_for('view_notifications'))

@app.route('/notifications/mark-all-read')
@login_required
def mark_all_notifications_read():
    """Mark all notifications as read"""
    db_session = get_session(engine)
    
    # Update all notifications for the current user
    db_session.query(Notification).filter_by(
        user_id=current_user.id,
        is_read=False
    ).update({Notification.is_read: True})
    
    db_session.commit()
    db_session.close()
    
    return redirect(url_for('view_notifications'))

# Chat Routes
@app.route('/chat')
@login_required
def chat_dashboard():
    """Chat dashboard showing available channels"""
    db_session = get_session(engine)
    
    # Get general chat channel
    general_channel = db_session.query(ChatChannel).filter_by(
        channel_type=ChatChannelType.GENERAL.value
    ).first()
    
    # Get user's stations for station chats
    user_stations = []
    
    # If user is admin, get all stations
    if current_user.is_admin:
        stations = db_session.query(Station).all()
        for station in stations:
            user_stations.append({
                'id': station.id,
                'name': station.name
            })
    else:
        # Get stations where user is a member
        stations_query = db_session.query(Station).join(
            StationUser, Station.id == StationUser.station_id
        ).filter(
            StationUser.user_id == current_user.id
        )
        
        for station in stations_query:
            user_stations.append({
                'id': station.id,
                'name': station.name
            })
    
    # Get station chat channels for user's stations
    station_channels = []
    for station in user_stations:
        channel = db_session.query(ChatChannel).filter_by(
            station_id=station['id'],
            channel_type=ChatChannelType.STATION.value
        ).first()
        
        if channel:
            # Get last message time and username for preview
            last_message = db_session.query(ChatMessage).filter_by(
                channel_id=channel.id
            ).order_by(desc(ChatMessage.created_at)).first()
            
            last_message_time = last_message.created_at if last_message else None
            last_message_sender = last_message.sender.username if last_message and not last_message.is_system_message else "System"
            
            station_channels.append({
                'id': channel.id,
                'name': channel.name,
                'station_id': station['id'],
                'station_name': station['name'],
                'last_message_time': last_message_time,
                'last_message_sender': last_message_sender
            })
    
    # Get last message for general channel
    last_general_message = None
    last_general_sender = None
    if general_channel:
        last_message = db_session.query(ChatMessage).filter_by(
            channel_id=general_channel.id
        ).order_by(desc(ChatMessage.created_at)).first()
        
        if last_message:
            last_general_message = last_message.created_at
            last_general_sender = last_message.sender.username if not last_message.is_system_message else "System"
    
    db_session.close()
    
    return render_template('chat_dashboard.html',
                          general_channel=general_channel,
                          station_channels=station_channels,
                          last_general_message=last_general_message,
                          last_general_sender=last_general_sender)

@app.route('/chat/channel/<int:channel_id>')
@login_required
def view_chat_channel(channel_id):
    """View a chat channel and its messages"""
    db_session = get_session(engine)
    
    # Get the channel
    channel = db_session.query(ChatChannel).get(channel_id)
    
    if not channel:
        db_session.close()
        flash('Chat channel not found', 'danger')
        return redirect(url_for('chat_dashboard'))
    
    # Check permissions for station channels
    if channel.channel_type == ChatChannelType.STATION.value:
        # If this is a station channel, check if user is in the station
        if not current_user.is_admin:
            station_user = db_session.query(StationUser).filter_by(
                station_id=channel.station_id,
                user_id=current_user.id
            ).first()
            
            if not station_user:
                db_session.close()
                flash('You do not have access to this chat channel', 'danger')
                return redirect(url_for('chat_dashboard'))
    
    # Get messages for this channel
    messages = db_session.query(ChatMessage).filter_by(
        channel_id=channel.id
    ).order_by(ChatMessage.created_at.desc()).limit(100).all()
    
    # Reverse messages to show oldest first
    messages.reverse()
    
    # Get user's available channels for the sidebar
    general_channel = db_session.query(ChatChannel).filter_by(
        channel_type=ChatChannelType.GENERAL.value
    ).first()
    
    user_stations = []
    if current_user.is_admin:
        # Get all stations for admin
        stations = db_session.query(Station).all()
        for station in stations:
            user_stations.append(station.id)
    else:
        # Get user's stations
        stations = db_session.query(Station).join(
            StationUser, Station.id == StationUser.station_id
        ).filter(
            StationUser.user_id == current_user.id
        )
        
        for station in stations:
            user_stations.append(station.id)
    
    # Get station chat channels for user's stations
    station_channels = db_session.query(ChatChannel).filter(
        ChatChannel.station_id.in_(user_stations),
        ChatChannel.channel_type == ChatChannelType.STATION.value
    ).all()
    
    db_session.close()
    
    return render_template('chat_channel.html',
                          channel=channel,
                          messages=messages,
                          general_channel=general_channel,
                          station_channels=station_channels,
                          current_channel_id=channel_id)

@app.route('/chat/channel/<int:channel_id>/send', methods=['POST'])
@login_required
def send_chat_message(channel_id):
    """Send a message to a chat channel"""
    db_session = get_session(engine)
    
    # Get the channel
    channel = db_session.query(ChatChannel).get(channel_id)
    
    if not channel:
        db_session.close()
        flash('Chat channel not found', 'danger')
        return redirect(url_for('chat_dashboard'))
    
    # Check permissions for station channels
    if channel.channel_type == ChatChannelType.STATION.value:
        # If this is a station channel, check if user is in the station
        if not current_user.is_admin:
            station_user = db_session.query(StationUser).filter_by(
                station_id=channel.station_id,
                user_id=current_user.id
            ).first()
            
            if not station_user:
                db_session.close()
                flash('You do not have access to this chat channel', 'danger')
                return redirect(url_for('chat_dashboard'))
    
    # Get message content from form
    content = request.form.get('message')
    
    if not content or content.strip() == '':
        db_session.close()
        flash('Message cannot be empty', 'danger')
        return redirect(url_for('view_chat_channel', channel_id=channel_id))
    
    # Create and save the message
    message = ChatMessage(
        channel_id=channel.id,
        sender_id=current_user.id,
        content=content
    )
    
    db_session.add(message)
    db_session.commit()
    db_session.close()
    
    return redirect(url_for('view_chat_channel', channel_id=channel_id))

@app.route('/chat/api/messages/<int:channel_id>')
@login_required
def get_chat_messages(channel_id):
    """API to get messages for a channel (for polling updates)"""
    db_session = get_session(engine)
    
    # Get the channel
    channel = db_session.query(ChatChannel).get(channel_id)
    
    if not channel:
        db_session.close()
        return jsonify({'error': 'Channel not found'}), 404
    
    # Check permissions for station channels
    if channel.channel_type == ChatChannelType.STATION.value:
        # If this is a station channel, check if user is in the station
        if not current_user.is_admin:
            station_user = db_session.query(StationUser).filter_by(
                station_id=channel.station_id,
                user_id=current_user.id
            ).first()
            
            if not station_user:
                db_session.close()
                return jsonify({'error': 'Access denied'}), 403
    
    # Get timestamp for last message shown (for polling)
    last_timestamp = request.args.get('last_timestamp')
    
    # Query for messages
    query = db_session.query(ChatMessage).filter(
        ChatMessage.channel_id == channel.id
    )
    
    # If last_timestamp provided, only get newer messages
    if last_timestamp:
        try:
            last_dt = datetime.datetime.fromisoformat(last_timestamp)
            query = query.filter(ChatMessage.created_at > last_dt)
        except ValueError:
            pass
    
    # Get messages ordered by time
    messages = query.order_by(ChatMessage.created_at).all()
    
    # Format messages for JSON response
    message_list = []
    for message in messages:
        message_list.append({
            'id': message.id,
            'sender_id': message.sender_id,
            'sender_name': message.sender.username if not message.is_system_message else "System",
            'content': message.content,
            'created_at': message.created_at.isoformat(),
            'is_system': message.is_system_message,
            'is_own': message.sender_id == current_user.id
        })
    
    db_session.close()
    
    return jsonify({'messages': message_list})

@app.context_processor
def inject_unread_notifications_count():
    """Inject unread notifications count into all templates"""
    if current_user.is_authenticated:
        db_session = get_session(engine)
        count = db_session.query(Notification).filter_by(
            user_id=current_user.id,
            is_read=False
        ).count()
        db_session.close()
        return {'unread_notifications_count': count}
    return {'unread_notifications_count': 0}

if __name__ == '__main__':
    app.run(debug=True, port=6789)
